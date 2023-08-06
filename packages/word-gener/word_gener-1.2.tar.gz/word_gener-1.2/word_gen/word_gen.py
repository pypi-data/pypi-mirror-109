import csv
import cv2
import docx
from docx import Document
from docxcompose.composer import Composer
import numpy as np
import os
import random
from random import randint
import string


class Word:
    files = []
    trash_files = []

    def  random_drawing(shapes:list):
        """
        This function generates a random picture from geometric shapes
        """
        document = Document()

        # Create a black image
        img = np.zeros((512, 512, 3), np.uint8)

        # Draw a diagonal blue line with thickness of 5 px
        if 'line' in shapes:
            cv2.line(img, (0, 0), (randint(1, 511), randint(1, 511)), (randint(1, 511), 0, 0), 5)

        if 'rectangle' in shapes:
            cv2.rectangle(img, (randint(1, 511), randint(1, 511)), (randint(1, 511), randint(1, 511)),
                            (randint(1, 511), randint(1, 511), randint(1, 511)), 3)
        if 'circle' in shapes:
            cv2.circle(img, (randint(1, 511), randint(1, 511)), 63, (0, 0, randint(1, 511)), -1)
        if 'ellipse' in shapes:
            cv2.ellipse(img, (randint(1, 511), randint(1, 511)), (randint(1, 511), randint(1, 511)), 0, 0, 180, 255, -1)

        cv2.imwrite('Image.jpg', img)
        doc = docx.Document()

        doc.add_paragraph('Random generated picture')
        doc.add_picture('Image.jpg', width=docx.shared.Cm(15))
        document.add_page_break()
        doc.save('test.docx')
        Word.files.append('test.docx')
        Word.trash_files.append('Image.jpg')

    def get_table():
        """
        This function generates a random fixed-sized table of numbers
        :return:
        """
        file = open("test1.txt", "w")

        my_array = np.random.rand(10, 4)
        np.savetxt('test1.txt', my_array, fmt='%4.6f', delimiter=' ')
        with open('test1.txt') as infile, open('test1.csv', 'w') as outfile:
            for line in infile:
                outfile.write(line.replace(' ', ','))
        doc = docx.Document()
        file.close()

        with open('test1.csv', newline='') as f:
            csv_reader = csv.reader(f)
            doc.add_page_break()
            doc.add_paragraph('Random generated table')

            csv_headers = next(csv_reader)
            csv_cols = len(csv_headers)

            table = doc.add_table(rows=2, cols=csv_cols)
            hdr_cells = table.rows[0].cells

            for i in range(csv_cols):
                hdr_cells[i].text = csv_headers[i]

            for row in csv_reader:
                row_cells = table.add_row().cells
                for i in range(csv_cols):
                    row_cells[i].text = row[i]
        doc.save("test1.docx")
        Word.files.append("test1.docx")
        Word.trash_files.append('test1.txt')
        Word.trash_files.append('test1.csv')

    def merge(path:str):
        """
        This function combines several word files into one single file at the specified path. Possibly can be used as a standalone helper library or smth idk ¯\_(ツ)_/¯
        :param path: selected save path
        :return:
        """
        new_document = Document()
        composer = Composer(new_document)
        for fn in Word.files:
            composer.append(Document(fn))
        composer.save(path)
        print('Merged file just saved to ', path, '!')

    def trash_clean():
        """
        This function cleans up intermediate files. In fact, they can remove this program too
        """
        Word.irl_trash = Word.trash_files + Word.files
        for f in Word.irl_trash:
            os.remove(f)

    def random_string(len: int):
        """
        This function generates a random set of letters of a given length
        :param length: length- number of letters of random text
        :return:
        """
        length = max(len, 0)
        letters = string.ascii_lowercase
        rand_string = ''.join(random.choice(letters) for i in range(length))
        file = open('test.txt', 'w')
        print("Random string of length", length, "is:", rand_string, file=file)
        file.close()

        file = open("test.txt", "r")
        data = file.read()
        doc = docx.Document()
        doc.add_page_break()
        doc.add_paragraph('Random generated string:')
        doc.add_paragraph(data)
        doc.save('test2.docx')
        file.close()
        Word.files.append("test2.docx")
        Word.trash_files.append('test.txt')
        return length